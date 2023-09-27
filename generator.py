from docx import Document
from datetime import datetime


def current_date_formatted():
    current_date = datetime.now()
    formatted_date = current_date.strftime('%d.%m.%Y')
    return formatted_date


def generate_initials(full_name):
    parts = full_name.split()
    initials = [name[0].upper() + '.' for name in parts[1:3]]
    return ''.join(initials) + ' ' + parts[0]


my_dict = {'contract_number': 'num123', 'parent_name': 'Фамилия Имя Отчество:',
           'children_name': 'ФИО Ребёнка:', 'parent_phone': '899999999:',
           'parent_email': 'mail@example.com', 'passport_series': '1488', 'passport_number': '228',
           'passport_issue_date': '2023-09-07', 'passport_place_of_issue': 'Кем выдан паспорт:',
           'registration_address': 'Адрес регистрации:', 'course_name': 'Основной'}


def create_dict(input_dict):
    output_dict = input_dict
    output_dict['contract_date'] = current_date_formatted()
    output_dict['parent_initials'] = generate_initials(input_dict['parent_name'])
    if input_dict['course_name'] == 'Основной':
        output_dict['course_start_date'] = '30.10.2023'
        output_dict['course_duration'] = '112 академических часов'
    if input_dict['course_name'] == 'Углублённый':
        output_dict['course_start_date'] = '03.10.2023'
        output_dict['course_duration'] = '64 академических часа'

    return output_dict


def create_new_file(input_dict):
    doc = Document('file.docx')
    for key, value in input_dict.items():
        for paragraph in doc.paragraphs:
            while key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, input_dict[key])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    while key in cell.text:
                        cell.text = cell.text.replace(key, input_dict[key])
    doc_name = str(input_dict['parent_name']) + '__' + str(input_dict['timestamp'])
    doc.save(f'{doc_name}.docx')
